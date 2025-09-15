"""
DOCX and XLSX extractors for Microsoft Office documents.
"""
import logging
from pathlib import Path
from typing import Union, List, Dict, Any, Optional, Tuple
import re

from src.extractors.base import BaseExtractor, ExtractorError
from src.models.types import PageParse, Block, Span
from src.utils.bbox import normalize_bbox

# Optional dependencies with graceful fallbacks
try:
    from docx import Document
    from docx.shared import Inches
    from docx.table import Table as DocxTable
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX extraction will be disabled.")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logging.warning("pandas not available. XLSX processing will be limited.")

try:
    import openpyxl
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl not available. XLSX extraction will be disabled.")


class DOCXExtractor(BaseExtractor):
    """
    DOCX extractor using python-docx for Word document processing.
    
    Extracts text, tables, and structure from DOCX files with proper
    formatting preservation and table structure recognition.
    """
    
    def __init__(self, extract_tables: bool = True, preserve_formatting: bool = True):
        """
        Initialize the DOCX extractor.
        
        Args:
            extract_tables: Whether to extract table structures
            preserve_formatting: Whether to preserve text formatting information
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCXExtractor")
        
        self.extract_tables = extract_tables
        self.preserve_formatting = preserve_formatting
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """
        Check if this extractor can handle the given file.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if file is a DOCX file, False otherwise
        """
        if not PYTHON_DOCX_AVAILABLE:
            return False
        
        path = Path(file_path)
        if path.suffix.lower() != '.docx':
            return False
        
        # Try to open the document to verify it's valid
        try:
            Document(str(file_path))
            return True
        except Exception:
            return False
    
    def get_page_count(self, file_path: Union[str, Path]) -> int:
        """
        Get the total number of pages in the DOCX document.
        
        Note: DOCX doesn't have explicit page breaks in the same way as PDFs.
        We treat the entire document as one logical page.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Always returns 1 for DOCX documents
        """
        return 1
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return ['docx']
    
    def parse_page(self, file_path: Union[str, Path], page_no: int) -> PageParse:
        """
        Parse the DOCX document (treated as a single page).
        
        Args:
            file_path: Path to the DOCX file
            page_no: Page number (must be 0 for DOCX)
            
        Returns:
            PageParse object containing extracted content and metadata
            
        Raises:
            ExtractorError: If parsing fails or page_no is not 0
        """
        if page_no != 0:
            raise ExtractorError(f"DOCX documents only have one page, requested page {page_no}", 
                               self.get_extractor_name(), str(file_path), page_no)
        
        try:
            # Load the document
            doc = Document(str(file_path))
            
            # Extract blocks from document elements
            blocks = self._extract_blocks_from_document(doc)
            
            # Estimate document dimensions (standard letter size)
            width, height = 612.0, 792.0  # Points (8.5" x 11")
            
            return PageParse(
                page_no=0,
                width=int(width),
                height=int(height),
                blocks=blocks,
                artifacts_removed=[]
            )
            
        except Exception as e:
            raise ExtractorError(f"Failed to parse DOCX document: {e}", self.get_extractor_name(), str(file_path), 0)
    
    def _extract_blocks_from_document(self, doc: Document) -> List[Block]:
        """
        Extract blocks from DOCX document elements.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of Block objects
        """
        blocks = []
        current_y = 0.0  # Track vertical position for bbox estimation
        
        for element in doc.element.body:
            try:
                if element.tag.endswith('p'):  # Paragraph
                    paragraph = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            paragraph = p
                            break
                    
                    if paragraph:
                        block = self._extract_paragraph_block(paragraph, current_y)
                        if block:
                            blocks.append(block)
                            current_y += 0.05  # Estimate vertical spacing
                
                elif element.tag.endswith('tbl'):  # Table
                    if self.extract_tables:
                        table = None
                        for t in doc.tables:
                            if t._element == element:
                                table = t
                                break
                        
                        if table:
                            block = self._extract_table_block(table, current_y)
                            if block:
                                blocks.append(block)
                                current_y += 0.15  # Tables take more vertical space
            
            except Exception as e:
                logging.warning(f"Failed to process DOCX element: {e}")
                continue
        
        return blocks
    
    def _extract_paragraph_block(self, paragraph, y_position: float) -> Optional[Block]:
        """
        Extract a block from a DOCX paragraph.
        
        Args:
            paragraph: python-docx Paragraph object
            y_position: Estimated vertical position
            
        Returns:
            Block object or None if paragraph is empty
        """
        text = paragraph.text.strip()
        if not text:
            return None
        
        # Determine block type
        block_type = self._classify_paragraph(paragraph, text)
        
        # Create spans from runs
        spans = self._create_spans_from_paragraph(paragraph, y_position)
        
        # Estimate bbox
        bbox = [0.1, y_position, 0.9, y_position + 0.04]  # Normalized coordinates
        
        # Extract formatting metadata
        formatting_meta = self._extract_paragraph_formatting(paragraph)
        
        block = Block(
            type=block_type,
            text=text,
            html=None,
            bbox=bbox,
            spans=spans,
            meta={
                "source": "docx_paragraph",
                "style": paragraph.style.name if paragraph.style else None,
                **formatting_meta
            }
        )
        
        return block
    
    def _extract_table_block(self, table: DocxTable, y_position: float) -> Optional[Block]:
        """
        Extract a block from a DOCX table.
        
        Args:
            table: python-docx Table object
            y_position: Estimated vertical position
            
        Returns:
            Block object containing table data
        """
        try:
            # Extract table data
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if not table_data:
                return None
            
            # Convert to HTML and text
            html_table = self._table_data_to_html(table_data)
            text_table = self._table_data_to_text(table_data)
            
            # Estimate table bbox
            table_height = len(table_data) * 0.02  # Estimate based on row count
            bbox = [0.1, y_position, 0.9, y_position + table_height]
            
            # Create span for table
            span = Span(
                text=text_table,
                bbox=bbox,
                rot=0.0,
                conf=1.0
            )
            
            block = Block(
                type="table",
                text=text_table,
                html=html_table,
                bbox=bbox,
                spans=[span],
                meta={
                    "source": "docx_table",
                    "rows": len(table_data),
                    "cols": len(table_data[0]) if table_data else 0
                }
            )
            
            return block
            
        except Exception as e:
            logging.warning(f"Failed to extract DOCX table: {e}")
            return None
    
    def _classify_paragraph(self, paragraph, text: str) -> str:
        """
        Classify paragraph type based on style and content.
        
        Args:
            paragraph: python-docx Paragraph object
            text: Paragraph text
            
        Returns:
            Block type string
        """
        # Check style name
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name or 'title' in style_name:
                return "heading"
        
        # Check text patterns
        text_lower = text.lower().strip()
        
        # Title block patterns
        title_patterns = [
            r'sheet\s+\d+', r'drawing\s+no', r'project\s+no',
            r'scale\s*:', r'date\s*:', r'drawn\s+by', r'checked\s+by'
        ]
        
        for pattern in title_patterns:
            if re.search(pattern, text_lower):
                return "titleblock"
        
        # List patterns
        if re.match(r'^\s*[\d\w][\.\)]\s', text):
            return "list"
        
        # Short text might be heading
        if len(text.strip()) < 100 and any(keyword in text_lower for keyword in 
                                          ['section', 'part', 'chapter', 'division']):
            return "heading"
        
        return "paragraph"
    
    def _create_spans_from_paragraph(self, paragraph, y_position: float) -> List[Span]:
        """
        Create spans from paragraph runs.
        
        Args:
            paragraph: python-docx Paragraph object
            y_position: Vertical position
            
        Returns:
            List of Span objects
        """
        spans = []
        x_position = 0.1  # Start position
        
        for run in paragraph.runs:
            run_text = run.text
            if not run_text:
                continue
            
            # Estimate span width based on text length
            span_width = min(len(run_text) * 0.01, 0.8)  # Rough estimate
            
            span = Span(
                text=run_text,
                bbox=[x_position, y_position, x_position + span_width, y_position + 0.04],
                rot=0.0,
                conf=1.0
            )
            
            spans.append(span)
            x_position += span_width
        
        # If no runs, create a single span for the entire paragraph
        if not spans and paragraph.text.strip():
            span = Span(
                text=paragraph.text.strip(),
                bbox=[0.1, y_position, 0.9, y_position + 0.04],
                rot=0.0,
                conf=1.0
            )
            spans.append(span)
        
        return spans
    
    def _extract_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """
        Extract formatting information from paragraph.
        
        Args:
            paragraph: python-docx Paragraph object
            
        Returns:
            Dictionary of formatting metadata
        """
        formatting = {}
        
        try:
            # Alignment
            if paragraph.alignment is not None:
                formatting["alignment"] = str(paragraph.alignment)
            
            # Style information
            if paragraph.style:
                formatting["style_name"] = paragraph.style.name
            
            # Check for bold/italic in runs
            has_bold = any(run.bold for run in paragraph.runs if run.bold)
            has_italic = any(run.italic for run in paragraph.runs if run.italic)
            
            if has_bold:
                formatting["has_bold"] = True
            if has_italic:
                formatting["has_italic"] = True
        
        except Exception as e:
            logging.warning(f"Failed to extract paragraph formatting: {e}")
        
        return formatting
    
    def _table_data_to_html(self, table_data: List[List[str]]) -> str:
        """Convert table data to HTML format."""
        if not table_data:
            return ""
        
        html = "<table>\n"
        
        for i, row in enumerate(table_data):
            tag = "th" if i == 0 else "td"
            html += "  <tr>\n"
            for cell in row:
                cell_text = str(cell).strip() if cell else ""
                html += f"    <{tag}>{self._escape_html(cell_text)}</{tag}>\n"
            html += "  </tr>\n"
        
        html += "</table>"
        return html
    
    def _table_data_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to plain text format."""
        if not table_data:
            return ""
        
        text_lines = []
        for row in table_data:
            row_text = " | ".join(str(cell).strip() if cell else "" for cell in row)
            text_lines.append(row_text)
        
        return "\n".join(text_lines)
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (text.replace("&", "&amp;")
                   .replace("<", "&lt;")
                   .replace(">", "&gt;")
                   .replace('"', "&quot;")
                   .replace("'", "&#x27;"))


class XLSXExtractor(BaseExtractor):
    """
    XLSX extractor using pandas and openpyxl for Excel spreadsheet processing.
    
    Processes each sheet as a separate page with proper table extraction
    and metadata preservation.
    """
    
    def __init__(self, 
                 include_empty_cells: bool = False,
                 max_rows: Optional[int] = None,
                 max_cols: Optional[int] = None):
        """
        Initialize the XLSX extractor.
        
        Args:
            include_empty_cells: Whether to include empty cells in extraction
            max_rows: Maximum number of rows to process per sheet
            max_cols: Maximum number of columns to process per sheet
        """
        if not PANDAS_AVAILABLE or not OPENPYXL_AVAILABLE:
            raise ImportError("pandas and openpyxl are required for XLSXExtractor")
        
        self.include_empty_cells = include_empty_cells
        self.max_rows = max_rows
        self.max_cols = max_cols
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """
        Check if this extractor can handle the given file.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if file is an XLSX file, False otherwise
        """
        if not PANDAS_AVAILABLE or not OPENPYXL_AVAILABLE:
            return False
        
        path = Path(file_path)
        if path.suffix.lower() not in ['.xlsx', '.xls']:
            return False
        
        # Try to open the workbook to verify it's valid
        try:
            load_workbook(str(file_path), read_only=True)
            return True
        except Exception:
            return False
    
    def get_page_count(self, file_path: Union[str, Path]) -> int:
        """
        Get the total number of sheets (pages) in the XLSX file.
        
        Args:
            file_path: Path to the XLSX file
            
        Returns:
            Number of sheets in the workbook
        """
        try:
            workbook = load_workbook(str(file_path), read_only=True)
            sheet_count = len(workbook.sheetnames)
            workbook.close()
            return sheet_count
        except Exception as e:
            logging.error(f"Failed to get sheet count from XLSX: {e}")
            return 1
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return ['xlsx', 'xls']
    
    def parse_page(self, file_path: Union[str, Path], page_no: int) -> PageParse:
        """
        Parse a specific sheet (page) from the XLSX file.
        
        Args:
            file_path: Path to the XLSX file
            page_no: Sheet number to parse (0-indexed)
            
        Returns:
            PageParse object containing extracted content and metadata
            
        Raises:
            ExtractorError: If parsing fails
        """
        try:
            # Load workbook and get sheet names
            workbook = load_workbook(str(file_path), read_only=True)
            sheet_names = workbook.sheetnames
            
            if page_no >= len(sheet_names):
                workbook.close()
                raise ExtractorError(f"Sheet {page_no} does not exist (total sheets: {len(sheet_names)})", 
                                   self.get_extractor_name(), str(file_path), page_no)
            
            sheet_name = sheet_names[page_no]
            workbook.close()
            
            # Read the specific sheet using pandas - first check actual column count
            try:
                # Read without column limit first to check actual columns
                temp_df = pd.read_excel(str(file_path), sheet_name=sheet_name, nrows=1)
                actual_cols = len(temp_df.columns)
                
                # Use smaller of max_cols or actual columns
                cols_to_use = min(self.max_cols, actual_cols) if self.max_cols else actual_cols
                
                df = pd.read_excel(str(file_path), sheet_name=sheet_name, 
                                 nrows=self.max_rows, usecols=range(cols_to_use))
            except Exception:
                # Fallback: read without column restrictions
                df = pd.read_excel(str(file_path), sheet_name=sheet_name, nrows=self.max_rows)
            
            # Process the dataframe into blocks
            blocks = self._extract_blocks_from_dataframe(df, sheet_name)
            
            # Estimate sheet dimensions
            width, height = 792.0, 612.0  # Landscape orientation for spreadsheets
            
            return PageParse(
                page_no=page_no,
                width=int(width),
                height=int(height),
                blocks=blocks,
                artifacts_removed=[]
            )
            
        except Exception as e:
            raise ExtractorError(f"Failed to parse XLSX sheet {page_no}: {e}", self.get_extractor_name(), str(file_path), page_no)
    
    def _extract_blocks_from_dataframe(self, df: pd.DataFrame, sheet_name: str) -> List[Block]:
        """
        Extract blocks from pandas DataFrame.
        
        Args:
            df: pandas DataFrame containing sheet data
            sheet_name: Name of the sheet
            
        Returns:
            List of Block objects
        """
        blocks = []
        
        if df.empty:
            return blocks
        
        # Clean the dataframe
        if not self.include_empty_cells:
            df = df.dropna(how='all').dropna(axis=1, how='all')
        
        if df.empty:
            return blocks
        
        # Convert DataFrame to table data
        table_data = []
        
        # Add column headers
        headers = [str(col) for col in df.columns]
        table_data.append(headers)
        
        # Add data rows
        for _, row in df.iterrows():
            row_data = []
            for col in df.columns:
                cell_value = row[col]
                if pd.isna(cell_value):
                    cell_text = "" if self.include_empty_cells else ""
                else:
                    cell_text = str(cell_value)
                row_data.append(cell_text)
            table_data.append(row_data)
        
        # Create table block
        if table_data and len(table_data) > 1:  # Must have at least header + 1 row
            html_table = self._table_data_to_html(table_data)
            text_table = self._table_data_to_text(table_data)
            
            # Estimate table bbox (full sheet)
            bbox = [0.05, 0.05, 0.95, 0.95]
            
            # Create span
            span = Span(
                text=text_table,
                bbox=bbox,
                rot=0.0,
                conf=1.0
            )
            
            # Create table block
            table_block = Block(
                type="table",
                text=text_table,
                html=html_table,
                bbox=bbox,
                spans=[span],
                meta={
                    "source": "xlsx_sheet",
                    "sheet_name": sheet_name,
                    "rows": len(table_data),
                    "cols": len(table_data[0]) if table_data else 0,
                    "original_shape": df.shape
                }
            )
            
            blocks.append(table_block)
        
        return blocks
    
    def _table_data_to_html(self, table_data: List[List[str]]) -> str:
        """Convert table data to HTML format."""
        if not table_data:
            return ""
        
        html = "<table>\n"
        
        # Header row
        if len(table_data) > 0:
            html += "  <thead>\n    <tr>\n"
            for cell in table_data[0]:
                cell_text = str(cell).strip() if cell else ""
                html += f"      <th>{self._escape_html(cell_text)}</th>\n"
            html += "    </tr>\n  </thead>\n"
        
        # Data rows
        if len(table_data) > 1:
            html += "  <tbody>\n"
            for row in table_data[1:]:
                html += "    <tr>\n"
                for cell in row:
                    cell_text = str(cell).strip() if cell else ""
                    html += f"      <td>{self._escape_html(cell_text)}</td>\n"
                html += "    </tr>\n"
            html += "  </tbody>\n"
        
        html += "</table>"
        return html
    
    def _table_data_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to plain text format."""
        if not table_data:
            return ""
        
        # Calculate column widths
        col_widths = []
        for row in table_data:
            for i, cell in enumerate(row):
                cell_text = str(cell).strip() if cell else ""
                if i >= len(col_widths):
                    col_widths.append(len(cell_text))
                else:
                    col_widths[i] = max(col_widths[i], len(cell_text))
        
        # Format table
        text_lines = []
        for i, row in enumerate(table_data):
            formatted_cells = []
            for j, cell in enumerate(row):
                cell_text = str(cell).strip() if cell else ""
                width = col_widths[j] if j < len(col_widths) else 10
                formatted_cells.append(cell_text.ljust(width))
            
            text_lines.append(" | ".join(formatted_cells))
            
            # Add separator after header
            if i == 0 and len(table_data) > 1:
                separator = " | ".join(["-" * width for width in col_widths])
                text_lines.append(separator)
        
        return "\n".join(text_lines)
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (text.replace("&", "&amp;")
                   .replace("<", "&lt;")
                   .replace(">", "&gt;")
                   .replace('"', "&quot;")
                   .replace("'", "&#x27;"))